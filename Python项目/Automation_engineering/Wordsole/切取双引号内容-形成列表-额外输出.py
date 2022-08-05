# -*- coding:utf-8 -*-
# @Author ： 秉烛爱好者
# @File : 切取双引号内容-形成列表-额外输出
# @Project : Automation_engineering
# @Time : 2022/8/211:32
# Github : https://github.com/SuperLucas/ArtOfTesting.github.io.git
# Gitee : https://gitee.com/SuperLucas/ArtOfTesting.github.io.git

# 在此处下方键入你的内容：

# 目标1：换行输出每个双引号间的文本如“好家伙好看；后后加括号。”-->好家伙好看；后后加括号。
# 目标2：将文本写入docx文件换行输出-->好家伙;l;好看；后后加括号。
# 目标3：将文本再次在；后换行，并在docx输出
'''
好家伙好看；
后后加括号。
'''
# 目标4：在原docx文件中实现


'''
好家伙好看；
后后加括号。
'''
# 目标5：将上述文本在源文件中加粗
'''
from docx import document
import re

doc = Document(r"E:\提取测试.docx")
restr = '“(?:[^”])*”'

for p in doc.paragraphs:
    matchRet = re.findall(restr, p.text)
    for r in matchRet:
        mid = int(len(restr))/2
        p.text = p.text.replace(r,r[1:-1])
    p = document.add_paragraph(r[1:-1])
    r = paragraph.add_run(r,r[1:-1])
    r.bold = True
    p = document.add_paragraph(r[1])
    p = document.add_paragraph(r[mid]-1)
    p = document.add_paragraph(r[-1])
doc.save(r'E:\提取测试_修正4.docx')
'''

from docx import Document
# from docx import Document
import re
from docx.shared import RGBColor, Pt, Cm
import os
import glob
doc = Document(r"E:\提取测试\提取测试.docx")
restr = '“(?:[^“])*”'

# for p in doc.paragraphs:
#     matchRet = re.findall(restr,p.text)
#     for r in matchRet:
#         mid = int(len(restr))/2
#         p.text = p.text.replace(r,r[1:-1])
#         run = p.text
#         run.font.bold = True
#         p = Document().add_paragraph('duanluo')
#     print("这一段的run个数是：", len(doc.paragraphs[0].runs))
#     runs = doc.paragraphs[0].runs
#     for r in runs:
#         print(runs.index(r), r.text)
# doc.save(r"E:\测试\提取测试-02.docx")
####
# for p in doc.paragraphs:
#     matchRet = re.findall(restr,p.text)
#     for r in matchRet:
#         mid = int(len(restr))/2
#         p.text = p.text.replace(r,r[1:-1])
#         for paragraph in doc.paragraphs:
#             for run in paragraph.runs:
#                 p.font.bold = True
#                 p.font.color.rgb = RGBColor(255, 0, 0)
#     print("这一段的run个数是：", len(doc.paragraphs[0].runs))
#     runs = doc.paragraphs[0].runs
#     for r in runs:
#         print(runs.index(r), r.text)
# doc.save(r"E:\测试\提取测试-03.docx")
# for p in doc.paragraphs:
#     matchRet = re.findall(restr,p.text)
#     print(matchRet)
#     for r in matchRet:
#         mid = int(len(restr))/2
#         p.text = p.text.replace(r,r[1:-1])
#         print(p)
#         for paragraph in doc.paragraphs:
#             for run in paragraph.runs:
#                 p.font.bold = True
#                 p.font.color.rgb = RGBColor(255, 0, 0)
#     print("这一段的run个数是：", len(doc.paragraphs[0].runs))
#     runs = doc.paragraphs[0].runs
#     for r in runs:
#         print(runs.index(r), r.text)
# doc.save(r"E:\测试\提取测试-03.docx")
for p in doc.paragraphs:
    matchRet = re.findall(restr,p.text)
    print(matchRet)
    for r in matchRet:
        mid = int(len(restr))/2
        p.text = p.text.replace(r,r[1:-1])
        print(p)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                p.font.bold = True
                p.font.color.rgb = RGBColor(255, 0, 0)
    print("这一段的run个数是：", len(doc.paragraphs[0].runs))
    runs = doc.paragraphs[0].runs
    for r in runs:
        print(runs.index(r), r.text)
doc.save(r"E:\测试\提取测试-03.docx")


