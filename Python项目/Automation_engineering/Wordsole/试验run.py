# -*- coding:utf-8 -*-
# @Author ： 秉烛爱好者
# @File : 试验run
# @Project : Automation_engineering
# @Time : 2022/8/511:24
# Github : https://github.com/SuperLucas/ArtOfTesting.github.io.git
# Gitee : https://gitee.com/SuperLucas/ArtOfTesting.github.io.git

# 在此处下方键入你的内容：
from docx import Document
doc = Document(r"E:\提取测试\提取测试.docx")  #上面这段话保存在1.docx中
print("这一段的run个数是：",len(doc.paragraphs[0].runs))
runs = doc.paragraphs[0].runs
for r in runs:
    print(runs.index(r),r.text)