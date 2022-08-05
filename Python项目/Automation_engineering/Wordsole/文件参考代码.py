# -*- coding:utf-8 -*-
# @Author ： 秉烛爱好者
# @File : 文件参考代码
# @Project : Automation_engineering
# @Time : 2022/8/511:39
# Github : https://github.com/SuperLucas/ArtOfTesting.github.io.git
# Gitee : https://gitee.com/SuperLucas/ArtOfTesting.github.io.git

# 在此处下方键入你的内容：
# from docx import Document
# from docx.shared
# import RGBColor, Pt, Cm
# import osimport glob
# # 此处更换创建文件夹的路径
# mkdir_path = r'C:\Users\xxx\new_dir'
# # 此处更换所有文件所在的位置
# file_path = r'C:\Users\xxx\'
# from docx import Documentwordfile = Document(path)
#
# if not os.path.exists(mkdir_path):
#     os.mkdir(mkdir_path)
#     for file in glob.glob(file_path + '/*.docx'):
#         pass
#     for file in glob.glob(file_path + '/*.docx'):
#         docx = Document(file)
#     for paragraph in docx.paragraphs:
#         for
#     run in paragraph.runs:
#     pass
# ...
# for run in paragraph.runs:
#     if'资金' in run.text: run.font.bold = True
#     # 加粗
#     run.font.color.rgb = RGBColor(255, 0, 255) # 设置字体颜色
#     # 最后切记保存
# docx.save(mkdir_path + '/' + os.path.basename(file))


####
from docx import Document
from docx.shared import RGBColor, Pt, Cm
import os
import glob
mkdir_path = 'E:\测试'
if not os.path.exists(mkdir_path):
    os.mkdir(mkdir_path)
keyword = '大风'
file_path = 'E:\提取测试'
for file in glob.glob(file_path + '\*.docx'):
    docx = Document(file)
    for paragraph in docx.paragraphs:
        for run in paragraph.runs:
            if keyword in run.text:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
    docx.save('E:\测试\提取测试_04.docx')
