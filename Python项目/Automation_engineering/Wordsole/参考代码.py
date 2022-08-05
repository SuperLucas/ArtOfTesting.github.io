# -*- coding:utf-8 -*-
# @Author ： 秉烛爱好者
# @File : 参考代码
# @Project : Automation_engineering
# @Time : 2022/8/511:33
# Github : https://github.com/SuperLucas/ArtOfTesting.github.io.git
# Gitee : https://gitee.com/SuperLucas/ArtOfTesting.github.io.git

# 在此处下方键入你的内容：
import docx, os

# 获取待处理的文件的路径
path = '待处理文件'  # 文件所在文件夹
files = [path + "\\" + i for i in os.listdir(path)]  # 获取文件夹下的文件名,并拼接完整路径

# 逐个提取文件，设置字体格式
for file in files:
    doc = docx.Document(file)
    for run in doc.paragraphs[0].runs:  # 总标题字体格式
        F_title(run)

    for para in doc.paragraphs[1:3]:  # 部门、姓名及日期字体格式
        for run in para.runs:
            F_name_dept(run)

    title1 = ["一、", "二、", "三、", "四、"]  # 标题一的唯一特征字符串
    title2 = ["1、", "2、", "3、", "4、"]  # 标题二的唯一特征字符串
    for para in doc.paragraphs[3:]:
        if any(i in para.text for i in title1):  # 若该段落是标题一，则应用标题一的字体格式
            for run in para.runs:
                F_title1(run)
        elif any(j in para.text for j in title2):  # 若该段落是标题二，则应用标题二的字体格式
            for run in para.runs:
                F_title2(run)
        else:
            for run in para.runs:  # 其余都应用正文的字体格式
                F_main(run)
    doc.save('已处理文件\\{}'.format(file.split("\\")[1]))
